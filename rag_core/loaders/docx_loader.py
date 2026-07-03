from pathlib import Path
from docx import Document as WordDocument  # type: ignore
from .document import Document

class DOCXLoader:
    def load(self,path:str) ->  Document:
        doc = WordDocument(path)

        text="\n".join(
            paragraph.text
            for paragraph in doc.paragraphs
        )

        return Document(
            source=path,
            content=text,
            metadata={
                "type":"docx",
                "paragraphs":len(doc.paragraphs)
            }

        )